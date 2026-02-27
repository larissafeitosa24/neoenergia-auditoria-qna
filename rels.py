
# -*- coding: utf-8 -*-
"""
Neoenergia • Consulta Relatórios de Auditoria (Executivo + Auditor)
- KPIs e rankings 100% pandas (determinístico)
- Chat híbrido: tenta responder por pandas (campos/contagens) e,
  quando for pergunta estratégica/complexa, envia um "data pack" (tabelas + fatos)
  para o OpenAI responder como um executivo de Auditoria Interna, sem inventar.

Requisitos:
  pip install streamlit pandas scikit-learn python-dateutil requests openai reportlab python-docx

Secrets (Streamlit Cloud):
  OPENAI_API_KEY = "..."
"""

import io
import re
import json
import datetime
import unicodedata
from typing import Optional, List, Dict, Tuple

import requests
import pandas as pd
import streamlit as st
from dateutil.parser import parse as date_parse
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from openai import OpenAI

# Exportações
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.lib.colors import HexColor


# ===========================================================
# 🔧 CONFIGURAÇÃO FIXA — ALTERE APENAS AQUI
# ===========================================================
HEAD_URL = "https://raw.githubusercontent.com/larissafeitosa24/neoenergia-auditoria-qna/main/relatorios.csv"
FIND_URL = "https://raw.githubusercontent.com/larissafeitosa24/neoenergia-auditoria-qna/main/constatacoes.csv"
LOGO_URL = "https://raw.githubusercontent.com/larissafeitosa24/neoenergia-auditoria-qna/main/neo_logo.png"

NEO_GREEN = "#7CC04B"
NEO_BLUE  = "#0060A9"
NEO_DARK  = "#014e87"

DEFAULT_MODEL = "gpt-4o-mini"  # pode trocar

st.set_page_config(
    page_title="Neoenergia • Consulta Relatórios de Auditoria",
    page_icon="📗",
    layout="wide"
)

CSS = f"""
<style>
html, body, [class*="css"] {{
  font-family: Segoe UI, SegoeUI, Helvetica, Arial, sans-serif;
}}
h1, h2, h3 {{ color: {NEO_BLUE}; }}
.stButton>button {{
  background-color: {NEO_BLUE};
  color: white; border-radius: 6px;
}}
.stButton>button:hover {{ background-color: {NEO_DARK}; }}
.card {{
  border: 1px solid rgba(255,255,255,0.12);
  border-radius: 10px; padding: 14px; background: rgba(255,255,255,0.03);
}}
.small-note {{ color:#9aa4b2; font-size:12px; }}
.source {{
  border-left: 4px solid {NEO_BLUE};
  background: #f0f7ff; padding: 8px; margin: 6px 0;
}}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)


# ===========================================================
# Utilitários
# ===========================================================
@st.cache_data(show_spinner=False, ttl=300)
def load_csv(url: str) -> pd.DataFrame:
    r = requests.get(url, timeout=60)
    r.raise_for_status()
    data = r.content.decode("utf-8", errors="ignore")
    return pd.read_csv(io.StringIO(data))


@st.cache_data(show_spinner=False, ttl=1800)
def load_logo(url: str) -> Optional[bytes]:
    try:
        r = requests.get(url, timeout=20)
        r.raise_for_status()
        return r.content
    except Exception:
        return None


def _normalize_col(s: str) -> str:
    s = str(s or "").strip()
    s = unicodedata.normalize("NFKD", s).encode("ASCII", "ignore").decode("ASCII")
    s = s.replace("-", "_")
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^a-zA-Z0-9_]", "", s)
    return s.lower()


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [_normalize_col(c) for c in df.columns]
    return df


def ensure_col(df: pd.DataFrame, new_name: str, candidates: List[str], default="") -> pd.DataFrame:
    df = df.copy()
    for c in candidates:
        if c in df.columns:
            df[new_name] = df[c]
            return df
    if new_name not in df.columns:
        df[new_name] = default
    return df


def pick_col(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            return c
    return None


def to_iso(v) -> str:
    try:
        return date_parse(str(v), dayfirst=True).date().isoformat()
    except Exception:
        return ""


def safe_dt(v) -> pd.Timestamp:
    return pd.to_datetime(v, errors="coerce", dayfirst=True)


def normalize_priority(x: str) -> str:
    """
    Normaliza classificação/criticidade para: Alta, Média, Baixa, (ou "Não informado")
    Aceita variações: alta/high/muito alto/critico/critical, media/medium, baixa/low.
    """
    s = str(x or "").strip().lower()
    if not s or s in {"nan", "none", "null"}:
        return "Não informado"
    if "crit" in s or "muito alt" in s or "very high" in s or "highest" in s:
        return "Alta"
    if "alt" in s or "high" in s:
        return "Alta"
    if "med" in s or "medium" in s:
        return "Média"
    if "baix" in s or "low" in s:
        return "Baixa"
    if "Muito Alto" in s or "muito alt" in s or "critical" in s:
        return "Alta"
    return "Não informado"


def is_closed_status(s: str) -> bool:
    t = str(s or "").lower()
    return bool(re.search(r"(encerr|fech|closed|conclu|implement|finaliz)", t))


def extract_filters(q: str) -> Dict[str, Optional[str]]:
    q_up = q.upper()

    aud = None
    m = re.search(r"\bAUD[-_\s]?\d+\b", q_up)
    if m:
        aud = m.group(0).replace("_", "-").replace(" ", "-")

    # finding_id pode ser numérico; aqui aceitamos:
    # - "finding 123" / "finding_id 123" / "constatacao 123" / "ID 123"
    finding = None
    m2 = re.search(r"\b(finding_id|finding|constatacao|constatação|id)\s*[:#-]?\s*([a-zA-Z0-9._-]+)\b", q_up)
    if m2:
        finding = m2.group(2)

    years = re.findall(r"\b(20\d{2})\b", q)
    years = [str(y) for y in years]

    return {"aud_code": aud, "finding_id": finding, "years": years}


# ===========================================================
# Mapeamento de campos (como você listou)
# ===========================================================
HEAD_FIELDS = {
    "titulo_do_trabalho": ["title", "compromissos_da_auditoria"],
    "codigo_auditoria": ["aud_code"],
    "tipo_relatorio": ["report_type"],
    "empresa": ["company"],
    "data_emissao": ["emission_date"],
    "objetivo": ["objetivo"],
    "escopo": ["escopo"],
    "alcance": ["alcance"],
    "risco_processo": ["risco_processo"],
    "classificacao_trabalho": ["classification"],
    "ano": ["ano"],
    "mes": ["mes"],
    "cronograma_inicio": ["cronograma_inicio"],
    "cronograma_final": ["cronograma_final"],
    "cronograma_draft": ["cronograma_draft"],
}

FIND_FIELDS = {
    "aud_code": ["aud_code", "id_do_trabalho"],
    "finding_id": ["finding_id"],
    "nome_constatacao": ["nome_da_constatao", "nome_da_constatacao", "finding_title"],
    "descricao_constatacao": ["constatao", "constatacao", "finding_text", "resposta"],
    "tipo_constatacao": ["tipo_de_constatao", "tipo_de_constatacao"],
    "risco_associado": ["associated_main_risk_category"],
    "causa_raiz": ["root_cause_analysis", "root_cause_comments"],
    "classificacao_recomendacao": ["classificao", "classification", "criticidade", "prioridade", "severity"],
    "status_constatacao": ["status_da_constatao", "status"],
    "proprietario_constatacao": ["proprietrio_da_constatao", "proprietario_da_constatacao"],
    "area_responsavel": ["organization_of_finding_response_owner"],  
    "plano_recomendacao": ["descrio_do_plano_de_recomendao", "descricao_do_plano_de_recomendacao", "recommendation"],
    "data_acordada_inicial": ["data_acordada_inicial"],
    "vencimento": ["data_acordada__vencimento", "data_acordada_vencimento", "due_date"],
    "data_aprovada_atual": ["data_acordada_aprovada_atualmente"],
    "encerramento": ["end_date"],
    "progresso_validado": ["progresso_validado"],
    "estado_trabalho": ["estado_del_trabajo"],
    "negocio_associado": ["negcio_associado", "negocio_associado", "tema", "compromissos_da_auditoria"],
    "ano": ["ano"],
    "mes": ["mes"],
}

# "Auditor" não veio explícito como coluna oficial nos exemplos.
# Aqui usamos um fallback comum: "proprietario_da_constatacao" (quem "abriu"/dono)
# Se no seu CSV existir "auditor" ou "audit_lead", inclua aqui.
AUDITOR_CANDIDATES = [
    "auditor", "audit_lead", "auditor_responsavel",
    "proprietrio_da_constatao", "proprietario_da_constatacao",
    "proprietario_da_constatacao", "owner"
]


# ===========================================================
# Preparação dos dados (HEAD e FIND)
# ===========================================================
def prep_data(df_h_raw: pd.DataFrame, df_f_raw: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    df_h = normalize_columns(df_h_raw)
    df_f = normalize_columns(df_f_raw)

    # HEAD
    if "aud_code" not in df_h.columns:
        raise ValueError("O CSV de relatórios (relatorios.csv) não contém a coluna 'aud_code'.")

    df_h["aud_code"] = df_h["aud_code"].astype(str).str.strip().str.upper()

    # ano/mês
    if "ano" in df_h.columns:
        df_h["ano"] = df_h["ano"].astype(str).str.strip()
    else:
        df_h["ano"] = ""

    if "mes" in df_h.columns:
        df_h["mes"] = df_h["mes"].astype(str).str.strip()
    else:
        df_h["mes"] = ""

    # datas cronograma
    for c in ["cronograma_inicio", "cronograma_draft", "cronograma_final", "emission_date"]:
        if c in df_h.columns:
            df_h[c] = df_h[c].fillna("").astype(str)

    # FINDINGS
    if "aud_code" not in df_f.columns:
        df_f = ensure_col(df_f, "aud_code", ["id_do_trabalho"], default="")
    df_f["aud_code"] = df_f["aud_code"].astype(str).str.strip().str.upper()

    df_f = ensure_col(df_f, "finding_id", FIND_FIELDS["finding_id"], default="")
    df_f["finding_id"] = df_f["finding_id"].astype(str).str.strip()

    # garantir colunas canônicas
    df_f = ensure_col(df_f, "area_responsavel", FIND_FIELDS["area_responsavel"], default="")
    df_f = ensure_col(df_f, "status", FIND_FIELDS["status_constatacao"], default="")
    df_f = ensure_col(df_f, "due_date", FIND_FIELDS["vencimento"], default="")
    df_f = ensure_col(df_f, "classif_raw", FIND_FIELDS["classificacao_recomendacao"], default="")

    # normalizações úteis
    df_f["area_responsavel"] = df_f["area_responsavel"].fillna("").astype(str).str.strip()
    df_f["status"] = df_f["status"].fillna("").astype(str).str.strip()
    df_f["due_date"] = df_f["due_date"].fillna("").astype(str).str.strip()
    df_f["classif_norm"] = df_f["classif_raw"].apply(normalize_priority)

    # ano/mês em findings
    if "ano" in df_f.columns:
        df_f["ano"] = df_f["ano"].astype(str).str.strip()
    else:
        df_f["ano"] = ""

    if "mes" in df_f.columns:
        df_f["mes"] = df_f["mes"].astype(str).str.strip()
    else:
        df_f["mes"] = ""

    # auditor (fallback)
    auditor_col = pick_col(df_f, AUDITOR_CANDIDATES)
    if auditor_col:
        df_f["auditor"] = df_f[auditor_col].fillna("").astype(str).str.strip()
    else:
        df_f["auditor"] = ""  # pode ficar vazio; o ranking vai avisar

    # texto de constatação e título
    df_f = ensure_col(df_f, "finding_title", FIND_FIELDS["nome_constatacao"], default="")
    df_f = ensure_col(df_f, "finding_text", FIND_FIELDS["descricao_constatacao"], default="")
    df_f["finding_title"] = df_f["finding_title"].fillna("").astype(str)
    df_f["finding_text"] = df_f["finding_text"].fillna("").astype(str)

    return df_h, df_f


def overdue_flag(df_f: pd.DataFrame) -> pd.Series:
    due = safe_dt(df_f["due_date"])
    today = pd.Timestamp(datetime.date.today())
    closed = df_f["status"].apply(is_closed_status)
    return due.notna() & (due < today) & (~closed)


# ===========================================================
# KPIs / Rankings (Executivo)
# ===========================================================
def compute_kpis(df_h: pd.DataFrame, df_f: pd.DataFrame, years_filter: List[str], aud_subset: Optional[set] = None):
    # aplica filtros
    h = df_h.copy()
    f = df_f.copy()

    if years_filter:
        h = h[h["ano"].astype(str).isin(years_filter)]
        f = f[f["ano"].astype(str).isin(years_filter)]

    if aud_subset:
        h = h[h["aud_code"].isin(aud_subset)]
        f = f[f["aud_code"].isin(aud_subset)]

    # work count
    qtd_trabalhos = int(h["aud_code"].nunique()) if not h.empty else 0

    # recommendations/findings count
    f_valid = f[f["finding_id"].astype(str).str.strip() != ""].copy()
    qtd_recs = int(f_valid["finding_id"].nunique()) if not f_valid.empty else 0

    # by classification
    by_class = (
        f_valid.groupby("classif_norm")["finding_id"].nunique().reindex(["Alta", "Média", "Baixa", "Não informado"]).fillna(0).astype(int)
    )

    # overdue
    if not f_valid.empty:
        od = overdue_flag(f_valid)
        qtd_atraso = int(f_valid.loc[od, "finding_id"].nunique())
    else:
        qtd_atraso = 0

    return qtd_trabalhos, qtd_recs, by_class.to_dict(), qtd_atraso, h, f_valid


def top10_critical_works(df_h_filt: pd.DataFrame, df_f_filt: pd.DataFrame) -> pd.DataFrame:
    if df_h_filt.empty:
        return pd.DataFrame()

    f = df_f_filt.copy()
    if f.empty:
        out = df_h_filt[["aud_code", "title", "ano"]].copy()
        out["qtd_recomendacoes"] = 0
        out["alta"] = 0
        out["media"] = 0
        out["baixa"] = 0
        out["em_atraso"] = 0
        out["score_criticidade"] = 0
        return out.head(10)

    f["is_overdue"] = overdue_flag(f).astype(int)
    piv = (
        f.groupby(["aud_code", "classif_norm"])["finding_id"]
         .nunique()
         .unstack(fill_value=0)
    )
    for c in ["Alta", "Média", "Baixa", "Não informado"]:
        if c not in piv.columns:
            piv[c] = 0

    agg = f.groupby("aud_code").agg(
        qtd_recomendacoes=("finding_id", "nunique"),
        em_atraso=("is_overdue", "sum"),
    )

    # score (simples e efetivo)
    # Alta pesa 3, Média 2, Baixa 1, atraso pesa forte
    score = (piv["Alta"] * 3) + (piv["Média"] * 2) + (piv["Baixa"] * 1) + (agg["em_atraso"] * 4)
    out = agg.join(piv[["Alta", "Média", "Baixa"]], how="left")
    out["score_criticidade"] = score

    # junta com HEAD (sem empresa)
    h = df_h_filt.copy()
    h = h[["aud_code", "title", "ano"]].drop_duplicates()
    out = h.merge(out.reset_index(), on="aud_code", how="left").fillna(0)
    out.rename(columns={"Alta": "alta", "Média": "media", "Baixa": "baixa"}, inplace=True)

    # ordena e seleciona
    out = out.sort_values(["score_criticidade", "qtd_recomendacoes", "em_atraso"], ascending=False)
    return out[["aud_code", "title", "qtd_recomendacoes", "alta", "media", "baixa", "em_atraso", "ano"]].head(10)


def top3_areas(df_f_filt: pd.DataFrame) -> pd.DataFrame:
    f = df_f_filt.copy()
    if f.empty:
        return pd.DataFrame()
    f["is_overdue"] = overdue_flag(f).astype(int)
    g = (
        f.groupby("area_responsavel")
         .agg(
             qtd_recomendacoes=("finding_id", "nunique"),
             alta=("classif_norm", lambda s: int((s == "Alta").sum())),
             em_atraso=("is_overdue", "sum"),
         )
         .reset_index()
    )
    g["area_responsavel"] = g["area_responsavel"].replace("", "Não informado")
    return g.sort_values(["qtd_recomendacoes", "em_atraso", "alta"], ascending=False).head(3)


def top3_auditores(df_f_filt: pd.DataFrame) -> pd.DataFrame:
    f = df_f_filt.copy()
    if f.empty:
        return pd.DataFrame()
    f["auditor"] = f["auditor"].replace("", "Não informado")
    g = (
        f.groupby("auditor")
         .agg(
             qtd_recomendacoes=("finding_id", "nunique"),
             alta=("classif_norm", lambda s: int((s == "Alta").sum())),
         )
         .reset_index()
    )
    return g.sort_values(["qtd_recomendacoes", "alta"], ascending=False).head(3)


# ===========================================================
# Similaridade histórica (para ajudar escopo)
# ===========================================================
def similar_audits_by_title(df_h: pd.DataFrame, aud_code: str, top_n=8) -> pd.DataFrame:
    if "title" not in df_h.columns:
        return pd.DataFrame()

    row = df_h[df_h["aud_code"] == aud_code]
    if row.empty:
        return pd.DataFrame()

    target = str(row.iloc[0]["title"] or "").strip()
    if not target:
        return pd.DataFrame()

    corpus = df_h["title"].astype(str).fillna("").tolist()
    vect = TfidfVectorizer(strip_accents="unicode", ngram_range=(1, 2))
    X = vect.fit_transform(corpus)
    qv = vect.transform([target])
    sim = cosine_similarity(qv, X).flatten()

    out = df_h.copy()
    out["sim"] = sim
    out = out[out["aud_code"] != aud_code]
    out = out.sort_values("sim", ascending=False).head(top_n)
    return out[["aud_code", "title", "ano", "company", "sim"]]


def vulnerability_table(df_f: pd.DataFrame) -> pd.DataFrame:
    f = df_f.copy()
    if f.empty:
        return pd.DataFrame()

    due = safe_dt(f["due_date"])
    today = pd.Timestamp(datetime.date.today())
    days_to_due = (due - today).dt.days

    w = f["classif_norm"].map({"Alta": 3, "Média": 2, "Baixa": 1, "Não informado": 1}).fillna(1).astype(int)
    od = overdue_flag(f).astype(int)
    openish = (~f["status"].apply(is_closed_status)).astype(int)

    score = (od * 100) + (openish * 10) + (w * 5) + (days_to_due.fillna(999).clip(-999, 999) * -1)

    out = pd.DataFrame({
        "aud_code": f["aud_code"],
        "finding_id": f["finding_id"],
        "area_responsavel": f["area_responsavel"].replace("", "Não informado"),
        "auditor": f["auditor"].replace("", "Não informado"),
        "classificacao": f["classif_norm"],
        "status": f["status"],
        "vencimento": f["due_date"],
        "dias_para_vencer": days_to_due,
        "em_atraso": od.astype(bool),
        "score_vulnerabilidade": score
    })

    # adiciona título curto da constatação (se existir)
    if "finding_title" in f.columns:
        out["finding_title"] = f["finding_title"].astype(str)

    out = out.sort_values("score_vulnerabilidade", ascending=False).head(20)
    return out


# ===========================================================
# Chat: Respostas determinísticas + OpenAI (estratégico)
# ===========================================================
def openai_client() -> Optional[OpenAI]:
    api_key = st.secrets.get("OPENAI_API_KEY", "").strip()
    if not api_key:
        return None
    return OpenAI(api_key=api_key)


EXEC_SYSTEM_PROMPT = """
Você é um executivo sênior de Auditoria Interna (Superintendente), com postura estratégica e analítica.
Você deve:
- Priorizar relevância executiva: riscos, atrasos, criticidade, tendências, áreas reincidentes, impacto e ações.
- Ajudar auditores: sugerir pontos de escopo e foco com base no histórico e evidências disponíveis.
- NÃO inventar fatos. Use SOMENTE os dados fornecidos no "DATA PACK".
- Quando algo não existir nos dados, diga claramente.
- Sempre que possível, fundamente em números e cite a origem (ex.: "KPI", "Top10", "Vulnerabilidade", "Histórico Similar").
- Responda em PT-BR, direto e claro.
Formato desejado:
1) Resposta objetiva
2) Evidências (bullet points)
3) Recomendações/Ações (bullet points)
Se a pergunta for sobre escopo, inclua um checklist prático de testes/linhas de investigação.
"""


def data_pack_for_llm(
    question: str,
    df_h_filt: pd.DataFrame,
    df_f_filt: pd.DataFrame,
    filters: Dict[str, Optional[str]]
) -> Dict:
    aud = filters.get("aud_code")
    years = filters.get("years") or []

    # KPIs e rankings do contexto filtrado
    qtd_trabalhos, qtd_recs, by_class, qtd_atraso, _, _ = compute_kpis(df_h_filt, df_f_filt, years_filter=years, aud_subset=None)

    top10 = top10_critical_works(df_h_filt, df_f_filt)
    top3a = top3_areas(df_f_filt)
    top3u = top3_auditores(df_f_filt)
    vuln = vulnerability_table(df_f_filt)

    pack = {
        "question": question,
        "filters_detected": filters,
        "kpis": {
            "qtd_trabalhos_aud_code": qtd_trabalhos,
            "qtd_recomendacoes_finding_id_unicos": qtd_recs,
            "recomendacoes_por_classificacao": by_class,
            "recomendacoes_em_atraso": qtd_atraso
        },
        "tables": {
            "top10_trabalhos_criticos": top10.to_dict(orient="records"),
            "top3_areas_com_mais_recomendacoes": top3a.to_dict(orient="records"),
            "top3_auditores_com_mais_recomendacoes": top3u.to_dict(orient="records"),
            "top20_planos_vulneraveis_ao_atraso": vuln.to_dict(orient="records"),
        }
    }

    # contexto específico do AUD, se fornecido
    if aud:
        head_row = df_h_filt[df_h_filt["aud_code"] == aud].head(1)
        if not head_row.empty:
            hr = head_row.iloc[0].to_dict()
            # enxuga para só campos relevantes
            keep = ["aud_code","title","report_type","company","emission_date","objetivo","escopo","alcance","risco_processo","classification","ano","mes",
                    "cronograma_inicio","cronograma_draft","cronograma_final"]
            pack["audit_head"] = {k: str(hr.get(k, "")) for k in keep if k in hr}

        # findings do aud (amostra)
        f_aud = df_f_filt[df_f_filt["aud_code"] == aud].copy()
        if not f_aud.empty:
            f_aud["em_atraso"] = overdue_flag(f_aud)
            f_aud = f_aud.sort_values(["em_atraso","classif_norm"], ascending=[False, True]).head(30)
            pack["audit_findings_sample"] = f_aud[[
                "finding_id","classif_norm","status","due_date","area_responsavel","auditor","finding_title","finding_text"
            ]].to_dict(orient="records")

        # histórico similar (título)
        sim = similar_audits_by_title(df_h_filt, aud, top_n=8)
        if not sim.empty:
            pack["similar_audits_by_title"] = sim.to_dict(orient="records")

    return pack


def llm_answer(question: str, pack: Dict) -> str:
    client = openai_client()
    if not client:
        return "❌ OPENAI_API_KEY não encontrada no Secrets."

    # Evita excesso (pack pode crescer). Mantém uma versão compacta em JSON.
    payload = json.dumps(pack, ensure_ascii=False)

    resp = client.responses.create(
        model=DEFAULT_MODEL,
        input=[
            {"role": "system", "content": EXEC_SYSTEM_PROMPT},
            {"role": "user", "content": f"DATA PACK (JSON):\n{payload}\n\nTarefa: responda à pergunta acima usando SOMENTE o DATA PACK."}
        ],
        temperature=0.2,
        store=False,
    )
    return (resp.output_text or "").strip()


def try_deterministic_answer(question: str, df_h: pd.DataFrame, df_f: pd.DataFrame) -> Optional[str]:
    """
    Responde perguntas diretas por pandas (sem LLM), para evitar “se perder”.
    Retorna None se não reconhecer intenção (aí vai pro LLM).
    """
    q = question.strip()
    qn = q.lower()
    f = extract_filters(q)
    aud = f.get("aud_code")
    finding = f.get("finding_id")
    years = f.get("years") or []

    # Helpers: filtrar
    h = df_h.copy()
    ff = df_f.copy()
    if years:
        h = h[h["ano"].astype(str).isin(years)]
        ff = ff[ff["ano"].astype(str).isin(years)]

    if aud:
        h_aud = h[h["aud_code"] == aud]
        f_aud = ff[ff["aud_code"] == aud]
    else:
        h_aud = pd.DataFrame()
        f_aud = pd.DataFrame()

    if finding:
        f_find = ff[ff["finding_id"].astype(str) == str(finding)]
    else:
        f_find = pd.DataFrame()

    # 1) “Quando vai começar os trabalhos?” / cronograma
    if any(t in qn for t in ["quando começa", "quando vai comecar", "início", "inicio", "cronograma"]):
        if not aud:
            return "Me diga o **AUD-xxxx** para eu retornar o cronograma (início/fim)."
        if h_aud.empty:
            return f"Não encontrei o aud_code {aud} no relatorios.csv."
        row = h_aud.iloc[0]
        ini = to_iso(row.get("cronograma_inicio", ""))
        fim = to_iso(row.get("cronograma_final", ""))
        dra = to_iso(row.get("cronograma_draft", ""))
        return (
            f"**{aud} — {row.get('title','')}**\n"
            f"- Início (cronograma_inicio): **{ini or 'não informado'}**\n"
            f"- Fim (cronograma_final): **{fim or 'não informado'}**\n"
            f"- Draft (cronograma_draft): **{dra or 'não informado'}**"
        )

    # 2) “Quantos trabalhos…” (aud_code)
    if ("quant" in qn and ("trabalh" in qn or "aud_code" in qn or "aud-code" in qn)) and ("recomend" not in qn and "constat" not in qn):
        n = int(h["aud_code"].nunique()) if not h.empty else 0
        if years:
            return f"Quantidade de trabalhos (aud_code únicos) em {', '.join(years)}: **{n}**."
        return f"Quantidade de trabalhos (aud_code únicos): **{n}**."

    # 3) “Quantas constatações/recomendações…”
    if "quant" in qn and ("recomend" in qn or "constat" in qn):
        d = ff[ff["finding_id"].astype(str).str.strip() != ""]
        if aud:
            d = d[d["aud_code"] == aud]
        n = int(d["finding_id"].nunique()) if not d.empty else 0
        base = f"Quantidade de recomendações/constatações (finding_id únicos)"
        if aud:
            base += f" em **{aud}**"
        if years:
            base += f" ({', '.join(years)})"
        return f"{base}: **{n}**."

    # 4) “Em atraso…”
    if "atras" in qn:
        d = ff[ff["finding_id"].astype(str).str.strip() != ""].copy()
        if aud:
            d = d[d["aud_code"] == aud]
        if d.empty:
            return "Não encontrei registros para calcular atraso com os filtros atuais."
        od = overdue_flag(d)
        n = int(d.loc[od, "finding_id"].nunique())
        return f"Recomendações em atraso (prazo vencido e não encerradas): **{n}**."

    # 5) campos diretos do HEAD
    direct_head = [
        ("titulo", "title"),
        ("empresa", "company"),
        ("objetivo", "objetivo"),
        ("escopo", "escopo"),
        ("alcance", "alcance"),
        ("risco", "risco_processo"),
        ("classifica", "classification"),
        ("tipo de rel", "report_type"),
        ("emitid", "emission_date"),
    ]
    if aud and not h_aud.empty:
        for key, col in direct_head:
            if key in qn and col in h_aud.columns:
                v = str(h_aud.iloc[0].get(col, "") or "").strip()
                return v if v else "Campo vazio para esse trabalho."

    # 6) campos diretos do FINDINGS (por finding_id ou por aud_code)
    if ("nome" in qn and "constat" in qn) or ("quais sao os nomes" in qn) or ("quais são os nomes" in qn):
        d = f_find if not f_find.empty else f_aud
        if d.empty:
            return "Me diga o **finding_id** ou o **AUD-xxxx** para listar os nomes das constatações."
        names = d["finding_title"].astype(str).str.strip()
        names = names[names != ""].unique().tolist()
        if not names:
            return "Não encontrei nomes de constatação nos registros filtrados."
        return "\n".join([f"- {n}" for n in names[:25]]) + ("" if len(names) <= 25 else f"\n(+{len(names)-25} itens)")

    if ("descrev" in qn and "constat" in qn) or ("descricao" in qn and "constat" in qn):
        if f_find.empty:
            return "Me diga o **finding_id** para eu descrever a constatação."
        txt = str(f_find.iloc[0].get("finding_text", "") or "").strip()
        return txt if txt else "Campo de descrição está vazio para esse finding_id."

    # Se não reconheceu, devolve None e o LLM assume (com data pack)
    return None


# ===========================================================
# Exportações (simples)
# ===========================================================
def export_pdf(text: str, logo_bytes: Optional[bytes]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4
    y = H - 50
    if logo_bytes:
        c.drawImage(ImageReader(io.BytesIO(logo_bytes)), 40, y - 40, width=150, height=40)
        y -= 60
    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(HexColor(NEO_BLUE))
    c.drawString(40, y, "Neoenergia — Q&A de Relatórios")
    y -= 24
    c.setFont("Helvetica", 10)
    c.setFillColor(HexColor("#111111"))
    for line in text.split("\n"):
        if y < 50:
            c.showPage()
            y = H - 50
        line = re.sub(r"\*\*|_", "", line)
        c.drawString(40, y, line[:1400])
        y -= 14
    c.save()
    return buf.getvalue()


def export_docx(text: str, logo_bytes: Optional[bytes]) -> bytes:
    doc = Document()
    if logo_bytes:
        try:
            doc.add_picture(io.BytesIO(logo_bytes), width=Inches(2.2))
        except Exception:
            pass
    doc.add_heading("Neoenergia — Q&A de Relatórios", level=1)
    for line in text.split("\n"):
        doc.add_paragraph(re.sub(r"\*\*|_", "", line))
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ===========================================================
# APP
# ===========================================================
with st.spinner("Carregando dados do GitHub..."):
    df_h_raw = load_csv(HEAD_URL)
    df_f_raw = load_csv(FIND_URL)

df_h, df_f = prep_data(df_h_raw, df_f_raw)

logo_bytes = load_logo(LOGO_URL)
if logo_bytes:
    st.image(logo_bytes, width=180)

st.title("📗 Consulta Relatórios de Auditoria")
st.caption("Visão executiva (KPIs + rankings) + Chat analítico/estratégico (pandas + OpenAI).")

# -------------------- filtros de contexto --------------------
st.subheader("🔎 Filtros de contexto")
c1, c2, c3, c4 = st.columns(4)

with c1:
    year_opts = sorted(df_h["ano"].replace("", "Sem ano").unique().tolist())
    f_year = st.multiselect("Ano (HEAD.ano)", year_opts)

with c2:
    # títulos (opcional)
    title_opts = sorted(pd.Series(df_h.get("title", pd.Series([], dtype=str))).fillna("").astype(str).unique().tolist())
    title_opts = [t for t in title_opts if t.strip()]
    f_title = st.multiselect("Título do trabalho", title_opts)

with c3:
    # área (do FIND)
    area_opts = sorted(df_f["area_responsavel"].replace("", "Não informado").unique().tolist())
    f_area = st.multiselect("Área responsável (organization_of_finding_response_owner)", area_opts)

with c4:
    # classificação (do FIND)
    class_opts = ["Alta", "Média", "Baixa", "Não informado"]
    f_class = st.multiselect("Classificação recomendação (normalizada)", class_opts)

# aplica filtros
df_h_filt = df_h.copy()
df_f_filt = df_f.copy()

if f_year:
    years_norm = [("" if y == "Sem ano" else str(y)) for y in f_year]
    df_h_filt = df_h_filt[df_h_filt["ano"].astype(str).isin(years_norm)]
    df_f_filt = df_f_filt[df_f_filt["ano"].astype(str).isin(years_norm)]

if f_title:
    if "title" in df_h_filt.columns:
        df_h_filt = df_h_filt[df_h_filt["title"].isin(f_title)]

aud_subset = set(df_h_filt["aud_code"].unique().tolist())
if aud_subset:
    df_f_filt = df_f_filt[df_f_filt["aud_code"].isin(aud_subset)]

if f_area:
    df_f_filt = df_f_filt[df_f_filt["area_responsavel"].replace("", "Não informado").isin(f_area)]

if f_class:
    df_f_filt = df_f_filt[df_f_filt["classif_norm"].isin(f_class)]

# -------------------- visão executiva --------------------
st.subheader("📌 Resumo executivo")

qtd_trabalhos, qtd_recs, by_class, qtd_atraso, _, _ = compute_kpis(
    df_h_filt, df_f_filt, years_filter=[], aud_subset=None
)

k1, k2, k3, k4 = st.columns(4)
k1.metric("Qtd de trabalhos (aud_code)", f"{qtd_trabalhos}")
k2.metric("Qtd de recomendações (finding_id únicos)", f"{qtd_recs}")
k3.metric("Recomendações em atraso", f"{qtd_atraso}")
k4.metric("Alta / Média / Baixa", f"{by_class.get('Alta',0)} / {by_class.get('Média',0)} / {by_class.get('Baixa',0)}")

st.caption("Classificação é normalizada a partir do campo de recomendação (ex.: Alta/Média/Baixa). Atraso = vencimento < hoje e status não encerrado.")

t10 = top10_critical_works(df_h_filt, df_f_filt)
t3a = top3_areas(df_f_filt)
t3u = top3_auditores(df_f_filt)

a, b = st.columns(2)
with a:
    st.markdown("**Top 10 trabalhos mais críticos (sem empresa)**")
    if t10.empty:
        st.info("Sem dados no filtro atual.")
    else:
        st.dataframe(t10, use_container_width=True)

with b:
    st.markdown("**Top 3 áreas com mais recomendações**")
    if t3a.empty:
        st.info("Sem dados no filtro atual.")
    else:
        st.dataframe(t3a, use_container_width=True)

st.markdown("**Top 3 'auditores' com mais recomendações**")
if t3u.empty:
    st.info("Não encontrei uma coluna clara de auditor (usei fallback). Se você me disser a coluna correta, eu ajusto e fica perfeito.")
else:
    st.dataframe(t3u, use_container_width=True)

with st.expander("📍 Top 20 planos mais vulneráveis ao atraso (para gestão/monitoramento)"):
    vt = vulnerability_table(df_f_filt)
    if vt.empty:
        st.info("Sem dados no filtro atual.")
    else:
        st.dataframe(vt, use_container_width=True)

# -------------------- chat (híbrido) --------------------
st.subheader("💬 Chat analítico/estratégico (pandas + OpenAI)")
st.caption("Dica: escreva AUD-xxxx quando quiser detalhes de um trabalho. Ex.: “AUD-1794361 quais riscos e o que priorizar no escopo?”")

show_pack = st.checkbox("Mostrar 'data pack' enviado ao modelo (debug)", value=False)

if "history" not in st.session_state:
    st.session_state["history"] = []

q = st.chat_input("Ex.: Quais trabalhos começam nos próximos 60 dias? / AUD-1794361 sugira escopo com base no histórico")

if q:
    # 1) tenta resposta determinística (pandas)
    det = try_deterministic_answer(q, df_h_filt, df_f_filt)

    if det is not None:
        answer = det
        pack_used = None
    else:
        # 2) se não for direta, monta um DATA PACK e chama o LLM (executivo)
        filters = extract_filters(q)
        pack = data_pack_for_llm(q, df_h_filt, df_f_filt, filters)
        answer = llm_answer(q, pack)
        pack_used = pack

    st.session_state["history"].append(("user", q))
    st.session_state["history"].append(("assistant", answer, pack_used))

# render histórico
for msg in st.session_state["history"]:
    if msg[0] == "user":
        with st.chat_message("user"):
            st.write(msg[1])
    else:
        with st.chat_message("assistant"):
            st.write(msg[1])
            if show_pack and msg[2] is not None:
                st.markdown("**DATA PACK (debug)**")
                st.json(msg[2])

# -------------------- export --------------------
st.subheader("📤 Exportar última resposta")
last_answer = None
for msg in reversed(st.session_state.get("history", [])):
    if msg[0] == "assistant":
        last_answer = msg[1]
        break

c1, c2 = st.columns(2)
with c1:
    if st.button("⬇️ Exportar PDF", disabled=(last_answer is None)):
        pdf = export_pdf(last_answer or "", logo_bytes)
        st.download_button("Baixar PDF", pdf, "neoenergia_qa.pdf", mime="application/pdf")

with c2:
    if st.button("⬇️ Exportar Word", disabled=(last_answer is None)):
        docx = export_docx(last_answer or "", logo_bytes)
        st.download_button("Baixar DOCX", docx, "neoenergia_qa.docx")


# -------------------- rodapé --------------------
st.markdown("<div class='small-note'>"
            "Observação: respostas diretas são calculadas por pandas. "
            "Perguntas estratégicas usam OpenAI com um pacote de dados (KPIs + rankings + amostras), sem inventar informações."
            "</div>", unsafe_allow_html=True)
